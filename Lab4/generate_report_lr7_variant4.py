import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


def main() -> None:
    lab4_dir = os.path.dirname(os.path.abspath(__file__))
    reg_img = os.path.join(lab4_dir, "LR_4_variant4_regression.png")
    int_img = os.path.join(lab4_dir, "LR_4_variant4_interpolation_P4.png")
    out_docx_ua = os.path.join(lab4_dir, "СіМШі_ЛР7_Коксюк_варіант4.docx")
    out_docx_ascii = os.path.join(lab4_dir, "LR7_report_variant4.docx")

    # Variant 4 data (Task 2)
    x = [2, 4, 6, 8, 10, 12]
    y = [6.5, 4.4, 3.8, 3.5, 3.1, 3.0]

    n = len(x)
    sx = sum(x)
    sy = sum(y)
    sxx = sum(v * v for v in x)
    sxy = sum(x[i] * y[i] for i in range(n))

    beta1 = (n * sxy - sx * sy) / (n * sxx - sx * sx)
    beta0 = (sy - beta1 * sx) / n

    # Interpolation data (Task 3) — from the methodological guide
    xi = [0.1, 0.3, 0.4, 0.6, 0.7]
    yi = [3.2, 3.0, 1.0, 1.8, 1.9]

    # Coefficients computed for P4(x) = a0 + a1*x + a2*x^2 + a3*x^3 + a4*x^4
    a0 = -8.18
    a1 = 186.25
    a2 = -864.0277777777778
    a3 = 1480.5555555555557
    a4 = -852.7777777777778

    p02_num, p02_den = 449, 90
    p05_num, p05_den = 319, 450
    p02 = p02_num / p02_den
    p05 = p05_num / p05_den

    doc = Document()

    # DSTU-like page setup (common academic standard):
    # left 30mm, right 10mm, top/bottom 20mm
    section = doc.sections[0]
    section.left_margin = Mm(30)
    section.right_margin = Mm(10)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)

    # Base font: Times New Roman 14
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)

    def add_center(text: str, bold: bool = False) -> None:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            p.runs[0].bold = bold

    # Title page (based on reference structure)
    add_center("МІНІСТЕРСТВО ОСВІТИ І НАУКИ УКРАЇНИ")
    add_center(
        "ДВНЗ «КИЇВСЬКИЙ НАЦІОНАЛЬНИЙ ЕКОМІЧНИЙ УНІВЕРСИТЕТ\nІМЕНІ ВАДИМА ГЕТЬМАНА»"
    )
    doc.add_paragraph("")
    add_center("Навчально-науковий інститут «Інститут інформаційних технологій в економіці»")
    add_center("Кафедра інформаційних систем в економіці")
    doc.add_paragraph("")
    doc.add_paragraph("")
    add_center("ЗВІТ", bold=True)
    add_center("Лабораторна робота №7")
    add_center("з дисципліни «Системи і методи штучного інтелекту»")
    add_center("Тема: Лінійна регресія. Метод найменших квадратів. Інтерполяція")
    add_center("Варіант 4")

    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("Виконав (-ла):")
    doc.add_paragraph("студент (-ка) 4 курсу, групи ІН-304")
    doc.add_paragraph("Коксюк Олег Віталійович")
    add_center("\nКиїв – 2026")

    doc.add_page_break()

    # 1. Theory
    p = doc.add_paragraph("1. ТЕОРЕТИЧНІ ВІДОМОСТІ")
    p.runs[0].bold = True
    doc.add_paragraph(
        "Регресія — це клас задач машинного навчання, у яких необхідно спрогнозувати числове "
        "значення (або вектор значень) на основі вхідних ознак. Лінійна регресія моделює "
        "залежність y від x лінійною функцією."
    )
    doc.add_paragraph(
        "Проста лінійна регресія задається рівнянням: y = β0 + β1·x + ε, де β0 — вільний член "
        "(перетин з віссю Y), β1 — кутовий коефіцієнт (нахил), ε — випадкова похибка."
    )
    doc.add_paragraph(
        "Метод найменших квадратів (МНК) полягає у знаходженні таких β0 та β1, які мінімізують "
        "суму квадратів відхилень між експериментальними значеннями yi та прогнозом ŷi = β0 + β1·xi:"
    )
    doc.add_paragraph("S(β0, β1) = Σ ( yi − (β0 + β1·xi) )² → min.")
    doc.add_paragraph("Оцінки параметрів за формулами МНК:")
    doc.add_paragraph(
        "β1 = ( n·Σ(xi·yi) − (Σxi)(Σyi) ) / ( n·Σ(xi²) − (Σxi)² )"
    )
    doc.add_paragraph("β0 = ( Σyi − β1·Σxi ) / n")
    doc.add_paragraph(
        "Інтерполяція — побудова функції, яка точно проходить через задані табличні точки. "
        "Для 5 точок будується інтерполяційний поліном степеня 4:"
    )
    doc.add_paragraph("P4(x) = a0 + a1·x + a2·x² + a3·x³ + a4·x⁴")

    # 2. Practical part
    p = doc.add_paragraph("\n2. ПРАКТИЧНА ЧАСТИНА")
    p.runs[0].bold = True

    p = doc.add_paragraph("2.1. Завдання 2 — лінійна регресія методом МНК (варіант 4)")
    p.runs[0].bold = True

    doc.add_paragraph("Експериментальні дані:")
    doc.add_paragraph(f"x = {x}")
    doc.add_paragraph(f"y = {y}")

    doc.add_paragraph("Обчислимо допоміжні суми:")
    doc.add_paragraph(f"n = {n}")
    doc.add_paragraph(f"Σx = {sx}")
    doc.add_paragraph(f"Σy = {sy:.1f}")
    doc.add_paragraph(f"Σx² = {sxx}")
    doc.add_paragraph(f"Σ(x·y) = {sxy:.1f}")

    doc.add_paragraph("Знайдемо параметри прямої регресії:")
    doc.add_paragraph(f"β1 = {beta1:.6f}")
    doc.add_paragraph(f"β0 = {beta0:.6f}")
    doc.add_paragraph(
        f"Отже, рівняння: ŷ = {beta0:.2f} + ({beta1:.2f})·x = {beta0:.2f} − {abs(beta1):.2f}·x"
    )

    doc.add_paragraph("Графік експериментальних точок та апроксимуючої прямої (МНК):")
    if os.path.exists(reg_img):
        doc.add_picture(reg_img, width=Mm(160))
    else:
        doc.add_paragraph(f"[Файл графіка не знайдено: {reg_img}]")

    p = doc.add_paragraph("\n2.2. Завдання 3 — інтерполяція поліномом степеню 4")
    p.runs[0].bold = True

    doc.add_paragraph("Табличні значення (5 точок):")
    doc.add_paragraph(f"x = {xi}")
    doc.add_paragraph(f"y = {yi}")

    doc.add_paragraph(
        "Складаємо систему X·A = Y, де X — матриця Вандермонда у вигляді [x⁰, x¹, x², x³, x⁴], "
        "A = (a0, a1, a2, a3, a4)ᵀ. Розв’язавши систему, отримуємо коефіцієнти полінома P4(x):"
    )
    doc.add_paragraph(f"a0 = {a0:.10f}")
    doc.add_paragraph(f"a1 = {a1:.10f}")
    doc.add_paragraph(f"a2 = {a2:.10f}")
    doc.add_paragraph(f"a3 = {a3:.10f}")
    doc.add_paragraph(f"a4 = {a4:.10f}")

    doc.add_paragraph("Графік інтерполяційного полінома та вихідних точок:")
    if os.path.exists(int_img):
        doc.add_picture(int_img, width=Mm(160))
    else:
        doc.add_paragraph(f"[Файл графіка не знайдено: {int_img}]")

    doc.add_paragraph("Значення функції у проміжних точках (за вимогою методички):")
    doc.add_paragraph(f"P4(0.2) = {p02_num}/{p02_den} ≈ {p02:.10f}")
    doc.add_paragraph(f"P4(0.5) = {p05_num}/{p05_den} ≈ {p05:.10f}")

    # Conclusion
    p = doc.add_paragraph("\nВИСНОВОК")
    p.runs[0].bold = True
    doc.add_paragraph(
        "У ході лабораторної роботи було опрацьовано поняття лінійної регресії та метод найменших "
        "квадратів. Для варіанта 4 за експериментальними даними побудовано лінійну модель "
        f"ŷ = {beta0:.2f} − {abs(beta1):.2f}·x та побудовано відповідний графік. Також виконано "
        "інтерполяцію таблично заданої функції поліномом 4-го степеня, побудовано графік "
        "інтерполяції та обчислено значення P4(0.2) і P4(0.5)."
    )

    # References
    p = doc.add_paragraph("\nСПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ")
    p.runs[0].bold = True
    doc.add_paragraph(
        "1. Методичні вказівки до лабораторної роботи №7 «Лінійна регресія. Метод найменших квадратів. Інтерполяція»."
    )
    doc.add_paragraph("2. Документація бібліотеки Matplotlib: https://matplotlib.org")
    doc.add_paragraph("3. Документація NumPy: https://numpy.org")

    doc.save(out_docx_ua)
    doc.save(out_docx_ascii)
    print(out_docx_ua)
    print(out_docx_ascii)


if __name__ == "__main__":
    main()

