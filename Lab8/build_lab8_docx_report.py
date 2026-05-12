# -*- coding: utf-8 -*-
"""
Генерує звіт ЛР-8 у форматі .docx за структурою референсу Lab2 (КНЕУ, ДСТУ-подібне оформлення).
Запуск: python build_lab8_docx_report.py
"""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Mm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


BASE = Path(__file__).resolve().parent
OUT_DOCX = BASE / "СіМШі_Лаб8_Коксюк.docx"

IMG_TF_LOSS = BASE / "outputs" / "01_loss_curve.png"
IMG_TF_KB = BASE / "outputs" / "02_kb_convergence.png"
IMG_TF_FIT = BASE / "outputs" / "03_data_and_fit.png"
IMG_K_MSE = BASE / "outputs_keras" / "keras_mse.png"
IMG_K_FIT = BASE / "outputs_keras" / "keras_fit.png"


def _set_run_font(run, name: str = "Times New Roman", size_pt: int = 14) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)
    rPr.append(rFonts)


def _configure_normal_style(doc: Document) -> None:
    sty = doc.styles["Normal"]
    sty.font.name = "Times New Roman"
    sty.font.size = Pt(14)
    pf = sty.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


def _page_setup(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_height = Mm(297)
    sec.page_width = Mm(210)
    sec.left_margin = Mm(30)
    sec.right_margin = Mm(15)
    sec.top_margin = Mm(20)
    sec.bottom_margin = Mm(20)


def _add_centered(doc: Document, text: str, bold: bool = False, size: int = 14) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    _set_run_font(run, size_pt=size)


def _add_heading_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    _set_run_font(run, size_pt=14)


def _add_body(doc: Document, text: str, indent_cm: float = 1.25) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Mm(indent_cm * 10)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    _set_run_font(run)


def _add_code_block(doc: Document, lines: str) -> None:
    for line in lines.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Mm(10)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(line)
        _set_run_font(run, size_pt=11)


def _add_figure(doc: Document, path: Path, caption: str, width_in: float = 5.8) -> None:
    if not path.is_file():
        p = doc.add_paragraph()
        p.add_run(f"[Файл не знайдено: {path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    _set_run_font(r, size_pt=12)


def _add_table_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    _set_run_font(run, size_pt=12)


def _table_metrics(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    hdr = ("Параметр / показник", "Значення", "Примітка")
    for i, h in enumerate(hdr):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for idx, (a, b, c) in enumerate(rows, start=1):
        table.rows[idx].cells[0].text = a
        table.rows[idx].cells[1].text = b
        table.rows[idx].cells[2].text = c


def build() -> Path:
    doc = Document()
    _configure_normal_style(doc)
    _page_setup(doc)

    # --- Титульний аркуш (за зразком Lab2) ---
    _add_centered(doc, "МІНІСТЕРСТВО ОСВІТИ І НАУКИ УКРАЇНИ", size=14)
    _add_centered(doc, "ДВНЗ «КИЇВСЬКИЙ НАЦІОНАЛЬНИЙ ЕКОНОМІЧНИЙ УНІВЕРСИТЕТ", size=14)
    _add_centered(doc, "ІМЕНІ ВАДИМА ГЕТЬМАНА»", size=14)
    _add_centered(doc, "Навчально-науковий інститут «Інститут інформаційних технологій в економіці»", size=14)
    _add_centered(doc, "Кафедра інформаційних систем в економіці", size=14)
    doc.add_paragraph()
    _add_heading_block(doc, "ЗВІТ")
    doc.add_paragraph()
    _add_centered(doc, "Лабораторна робота №8", bold=True)
    _add_centered(doc, 'з дисципліни «Системи і методи штучного інтелекту»')
    doc.add_paragraph()
    _add_centered(doc, "Виконав:")
    _add_centered(doc, "студент 4 курсу, групи ІН-304")
    _add_centered(doc, "Коксюк Олег Віталійович")
    doc.add_paragraph()
    _add_centered(doc, "Київ – 2026")

    doc.add_page_break()

    # --- Мета ---
    _heading_section(doc, "МЕТА ТА ЗАВДАННЯ РОБОТИ")
    _add_body(
        doc,
        "Метою лабораторної роботи є засвоєння базових прийомів побудови та навчання "
        "нейромережевих моделей засобами TensorFlow та високорівневого API Keras на прикладі "
        "лінійної регресії; закріплення понять обчислювального графа, змінних, функції втрат "
        "та стохастичного градієнтного спуску (SGD).",
    )
    _add_body(
        doc,
        "Для досягнення мети вирішувались такі завдання:",
    )
    for item in (
        "згенерувати синтетичні дані для лінійної залежності з гаусовим шумом та навчити параметри моделі ŷ = kx + b у стилі TensorFlow 1.x (tf.compat.v1: placeholder, Session, sess.run);",
        "проаналізувати динаміку функції втрат та збіжність коефіцієнтів k і b до істинних значень;",
        "побудувати графіки: динаміка loss, траєкторії параметрів, дані та навчена пряма;",
        "реалізувати ту саму постановку задачі через keras.Sequential та Dense(1), порівняти результати з нижньорівневою реалізацією.",
    ):
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.left_indent = Mm(12)
        p.paragraph_format.first_line_indent = Mm(-6)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run("– " + item)
        _set_run_font(run)

    _add_body(
        doc,
        "Методичне забезпечення: файл «СШІ_ЛР_8_НАВЧАННЯ_НЕЙРОМЕРЕЖ_TENSORFLOW.pdf». "
        "Код експериментів зібрано у каталозі Lab8; архів Lab8.zip містить узгоджені скрипти та звітні матеріали.",
    )

    # --- Завдання 8.1 ---
    doc.add_paragraph()
    _heading_section(doc, "ЗАВДАННЯ 8.1. ЛІНІЙНА РЕГРЕСІЯ У TENSORFLOW (TF.COMPAT.V1)")
    _sub_heading(doc, "8.1.1. Постановка задачі та теоретичні відомості")
    _add_body(
        doc,
        "Розглядається синтетична регресія: незалежна змінна x рівномірно розподілена на відрізку [0; 1], "
        "залежна змінна y = 2x + 1 + ε, де ε ~ N(0, σ²), σ = √2 (дисперсія шуму дорівнює 2). "
        "Навчається модель ŷ = kx + b з двома параметрами. Мінімізується середній квадрат помилки на міні-батчі "
        "(у методичних матеріалах також зустрічається сума квадратів по батчу — reduce_sum); "
        "для стабільності градієнтів у робочій програмі використано tf.reduce_mean, що відповідає тій самій задачі "
        "мінімізації з іншим масштабом градієнта та узгодженим кроком навчання.",
    )
    _add_body(
        doc,
        "Оптимізація: GradientDescentOptimizer зі швидкістю навчання 0,08; розмір міні-батча 100; кількість ітерацій 20 000. "
        "Початкові значення k та b ініціалізовані нулями.",
    )

    _sub_heading(doc, "8.1.2. Реалізація алгоритму")
    _add_body(
        doc,
        "Обчислювальний граф містить placeholder X форми [None, 1], placeholder y форми [None], "
        "навчальні змінні k та b, прогноз y_hat = squeeze(k·X + b). Навчальний цикл виконує sess.run "
        "для операції оптимізації та оновлення параметрів на випадкових підвибірках. Фрагмент ключової логіки:",
    )
    _add_code_block(
        doc,
        """graph = tf.Graph()
with graph.as_default():
    X = tf.compat.v1.placeholder(tf.float32, shape=[None, 1])
    y = tf.compat.v1.placeholder(tf.float32, shape=[None])
    k = tf.compat.v1.get_variable("k", shape=(), initializer=tf.zeros_initializer())
    b = tf.compat.v1.get_variable("b", shape=(), initializer=tf.zeros_initializer())
    y_hat = tf.squeeze(k * X + b, axis=[-1])
    loss = tf.reduce_mean(tf.square(y - y_hat))
    optimizer = tf.compat.v1.train.GradientDescentOptimizer(0.08).minimize(loss)""",
    )

    _sub_heading(doc, "8.1.3. Результати навчання та графічний аналіз")
    summary_path = BASE / "outputs" / "training_summary.txt"
    k_fin = b_fin = loss_fin = "—"
    if summary_path.is_file():
        txt = summary_path.read_text(encoding="utf-8")
        for line in txt.splitlines():
            if line.startswith("final_k="):
                k_fin = line.split("=", 1)[1].strip()
            elif line.startswith("final_b="):
                b_fin = line.split("=", 1)[1].strip()
            elif line.startswith("final_batch_loss="):
                loss_fin = line.split("=", 1)[1].strip()

    _add_table_caption(doc, "Таблиця 8.1 – Результати навчання моделі в TensorFlow")
    _table_metrics(
        doc,
        [
            ("Оцінка k", k_fin, "істинне значення 2,0"),
            ("Оцінка b", b_fin, "істинне значення 1,0"),
            ("Loss на останньому батчі", loss_fin, "MSE по міні-батчі (не монотонна через SGD)"),
            ("Кількість ітерацій", "20 000", "batch_size = 100"),
        ],
    )
    doc.add_paragraph()

    _add_body(
        doc,
        "На рисунку 8.1 показано зміну функції втрат по ітераціях. Коливання після початкового спаду є типовими для "
        "SGD: на кожному кроці оптимізується втрата на новій випадковій підвибірці.",
    )
    _add_figure(
        doc,
        IMG_TF_LOSS,
        "Рисунок 8.1 – Динаміка функції втрат (MSE на міні-батчі) під час навчання",
    )

    _add_body(
        doc,
        "Рисунок 8.2 ілюструє збіжність параметрів k та b до околу істинних значень; пунктирні лінії відповідають k = 2 та b = 1.",
    )
    _add_figure(
        doc,
        IMG_TF_KB,
        "Рисунок 8.2 – Зміна параметрів k та b у процесі навчання та істинні значення",
    )

    _add_body(
        doc,
        "Рисунок 8.3 демонструє відповідність навченої прямої розкиду точок та близькість до теоретичної прямої y = 2x + 1.",
    )
    _add_figure(
        doc,
        IMG_TF_FIT,
        "Рисунок 8.3 – Синтетичні дані, навчена пряма та істинна залежність без шуму",
    )

    _add_body(
        doc,
        "Висновки до завдання 8.1: нижньорівневий API TensorFlow дозволяє явно задати граф обчислень і цикл "
        "навчання з sess.run; отримані оцінки k та b узгоджені з істинними коефіцієнтами з урахуванням стохастичного шуму "
        "та обмеженого обсягу вибірки. Використання reduce_mean замість reduce_sum потребує відповідного підбору "
        "швидкості навчання та має бути відображене у звіті як методично обґрунтоване допущення.",
        indent_cm=1.25,
    )

    # --- Завдання 8.2 ---
    doc.add_paragraph()
    _heading_section(doc, "ЗАВДАННЯ 8.2. ЛІНІЙНА РЕГРЕСІЯ В KERAS (SEQUENTIAL + DENSE)")
    _sub_heading(doc, "8.2.1. Опис моделі та параметри навчання")
    _add_body(
        doc,
        "Для порівняння з реалізацією через граф TensorFlow побудовано модель keras.Sequential з шарами Input(shape=(1,)) "
        "та Dense(1, use_bias=True). Функція втрат – середньоквадратична помилка (mse), оптимізатор – SGD з learning_rate = 0,2. "
        "Навчання: 400 епох, batch_size = 100; початкові ваги та зміщення ініціалізовані нулями. Чисельні результати залежать від "
        "seed генератора випадкових чисел і можуть відрізнятися від завдання 8.1 навіть при тій самій постановці.",
    )

    _sub_heading(doc, "8.2.2. Результати та графіки")
    _add_body(
        doc,
        "На рисунку 8.4 наведено зменшення MSE по епохах під час виклику model.fit (за одну епоху виконується повний прохід по даних).",
    )
    _add_figure(
        doc,
        IMG_K_MSE,
        "Рисунок 8.4 – Динаміка MSE при навчанні моделі Keras",
    )

    _add_body(
        doc,
        "Рисунок 8.5 містить візуальне порівняння прогнозу Keras з істинною прямою на інтервалі x ∈ [0; 1].",
    )
    _add_figure(
        doc,
        IMG_K_FIT,
        "Рисунок 8.5 – Дані, навчена пряма Keras та істинна залежність y = 2x + 1",
    )

    _add_body(
        doc,
        "Висновки до завдання 8.2: високорівневий API Keras скорочує обсяг коду порівняно з явним циклом sess.run і забезпечує "
        "зручне логування історії навчання (history.history['loss']). Отримана пряма узгоджується з лінійним трендом даних; "
        "відмінності чисельних k, b від результатів завдання 8.1 пояснюються іншим seed та режимом оновлення ваг (епохи проти ітерацій SGD).",
        indent_cm=1.25,
    )

    # --- Загальні висновки ---
    doc.add_paragraph()
    _heading_section(doc, "ВИСНОВКИ")
    bullets = (
        "Опановано базовий робочий цикл TensorFlow 2 з використанням сумісного режиму tf.compat.v1 для моделювання навчання лінійної регресії.",
        "Підтверджено здатність SGD знаходити параметри лінійної моделі за зашумлених даних; графіки loss та параметрів відображають стохастичну природу міні-батч навчання.",
        "Модель Keras з одним шаром Dense еквівалентна лінійній регресії та дає результати, узгоджені з постановкою задачі.",
        "Для оформлення звіту за ДСТУ доцільно зберігати підписи таблиць над таблицею, підписи рисунків – під рисунками, вирівнювання по ширині сторінки – за вимогами методичних рекомендацій кафедри.",
    )
    for b in bullets:
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.left_indent = Mm(12)
        p.paragraph_format.first_line_indent = Mm(-6)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run("– " + b)
        _set_run_font(run)

    _add_body(
        doc,
        "Програмні файли: LR_8_tensorflow_linear_regression.py, LR_8_keras_linear_regression.py; графіки збережено у підкаталогах outputs та outputs_keras.",
    )

    doc.add_paragraph()
    _heading_section(doc, "СПИСОК ДЖЕРЕЛ")
    refs = (
        "Методичні вказівки до лабораторної роботи №8 з дисципліни «Системи і методи штучного інтелекту» (TensorFlow / Keras).",
        "Abadi M. et al. TensorFlow: Large-scale machine learning on heterogeneous systems [Електронний ресурс]. – Режим доступу: https://www.tensorflow.org/",
        "Chollet F. Deep Learning with Python. – Manning Publications, 2021.",
    )
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Mm(12.5)
        p.paragraph_format.left_indent = Mm(0)
        run = p.add_run(f"{i}. {r}")
        _set_run_font(run)

    doc.save(OUT_DOCX)
    return OUT_DOCX


def _heading_section(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    _set_run_font(run, size_pt=14)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def _sub_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    _set_run_font(run, size_pt=14)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)


if __name__ == "__main__":
    path = build()
    print(f"Збережено: {path}")
